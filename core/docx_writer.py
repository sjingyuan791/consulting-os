from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import re

class DocxWriter:
    def __init__(self):
        self.document = Document()
        self._setup_styles()

    def _setup_styles(self):
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Yu Mincho' # Japanese standard
        font.size = Pt(10.5)

    def add_title(self, text):
        self.document.add_heading(text, 0)

    def add_chapter(self, text):
        self.document.add_heading(text, 1)
        
    def add_section(self, text):
        self.document.add_heading(text, 2)

    def add_content_from_markdown(self, md_text):
        """
        Parses limited markdown (bold, bullet points) and adds to docx.
        """
        lines = md_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Bullet points
            if line.startswith('- ') or line.startswith('* '):
                p = self.document.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('### '):
                 self.document.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                 self.document.add_heading(line[3:], level=2)
            else:
                p = self.document.add_paragraph(line)
                
    def save(self, path):
        self.document.save(path)
